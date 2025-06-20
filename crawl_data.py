import requests
from bs4 import BeautifulSoup
from docx import Document
import os
import time
import random
from concurrent.futures import ThreadPoolExecutor, as_completed
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# =================== CẤU HÌNH ===================
output_dir = 'D:/crawl'
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'thu_tuc_hanh_chinh.docx')
TOTAL_PAGES = 9
MAX_THREADS = 6
ERROR_LOG = os.path.join(output_dir, 'error_urls.txt')

# =================== SESSION VỚI RETRY ===================
def create_retry_session():
    session = requests.Session()
    retry = Retry(
        total=5,
        backoff_factor=1,
        status_forcelist=[429, 500, 502, 503, 504],
        raise_on_status=False
    )
    adapter = HTTPAdapter(max_retries=retry)
    session.mount('http://', adapter)
    session.mount('https://', adapter)
    return session

session = create_retry_session()

# =================== TÁCH BỐ CỤC ===================
def format_detail_text(text):
    import re
    headers = [
        "Mã thủ tục", "Tên thủ tục", "Cấp thực hiện", "Lĩnh vực", "Thông tin công bố",
        "Cách thức nộp trực tuyến", "Thời hạn giải quyết", "Mức trực tuyến", "Lệ phí",
        "Phí", "Cơ quan thực hiện", "Đối tượng thực hiện", "Cách thức thực hiện",
        "Điều kiện thực hiện", "Số bộ hồ sơ", "Kết quả thực hiện", "Địa chỉ tiếp nhận hồ sơ",
        "Trình tự thực hiện", "Thành phần hồ sơ", "Căn cứ pháp lý", "Tình trạng hiệu lực"
    ]
    output = ""
    for header in headers:
        pattern = rf"({header})(.*?)(?=" + "|".join(headers) + "|$)"
        match = re.search(pattern, text, re.DOTALL)
        if match:
            title = match.group(1).strip()
            content = match.group(2).strip()
            output += f"\n\n🔹 **{title}**\n{content}"
    return output.strip()

# =================== LẤY CHI TIẾT TỪ LINK ===================
def get_thu_tuc_detail(url):
    try:
        print(f"Đang lấy chi tiết: {url}")
        response = session.get(url, timeout=30)
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')

            title_tag = soup.find('div', class_='box-tthc').find('h2', class_='tieude-h2')
            title = title_tag.get_text(strip=True) if title_tag else "Không có tiêu đề"

            detail_div = soup.find('div', class_='table-reponsive')
            if detail_div:
                raw_text = detail_div.get_text(separator=' ', strip=True)
                formatted_detail = format_detail_text(raw_text)

                #  Tìm link file Word
                word_links = []
                for a in detail_div.find_all('a', href=True):
                    href = a['href']
                    if '.doc' in href or '.docx' in href or 'download_file.jsp' in href:
                        file_url = href if href.startswith('http') else 'https://csdl.dichvucong.gov.vn' + href
                        file_name = a.get_text(strip=True)
                        word_links.append(f"{file_name}: {file_url}")

                # Thêm phần tài liệu nếu có
                if word_links:
                    formatted_detail += "\n\n **Tài liệu đính kèm**"
                    for link in word_links:
                        formatted_detail += f"\n📎 {link}"

                return {
                    'title': title,
                    'content': formatted_detail
                }
            else:
                return {
                    'title': title,
                    'content': "Không tìm thấy nội dung chi tiết."
                }
        else:
            print(f"Lỗi mã trạng thái: {response.status_code}")
    except Exception as e:
        print(f"Lỗi truy cập {url}: {e}")
        with open(ERROR_LOG, "a", encoding="utf-8") as f:
            f.write(url + "\n")
    return None

# =================== CRAWL DANH SÁCH LINK ===================
def crawl_detail_links(links):
    results = []
    with ThreadPoolExecutor(max_workers=MAX_THREADS) as executor:
        futures = [executor.submit(get_thu_tuc_detail, link) for link in links]
        for future in as_completed(futures):
            detail = future.result()
            if detail:
                results.append(detail)
            time.sleep(random.uniform(0.5, 1.0))  # Giảm tải server
    return results

# =================== CRAWL TỪ TRANG ===================
def crawl_page(page_number):
    print(f"\nĐang crawl trang {page_number}...")
    url = f"https://dichvucong.danang.gov.vn/thu-tuc-hanh-chinh?p_p_id=thutuchanhchinh_WAR_dngdvcportlet&p_p_lifecycle=0&p_p_state=normal&p_p_mode=view&p_p_col_id=column-2&p_p_col_count=1&_thutuchanhchinh_WAR_dngdvcportlet_delta=20&_thutuchanhchinh_WAR_dngdvcportlet_cur={page_number}"
    
    try:
        response = session.get(url, timeout=30)
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')
            tds = soup.find_all('td', class_='table-cell dvc-tendvc')
            links = []

            for td in tds:
                a_tag = td.find('a', href=True)
                if a_tag:
                    href = a_tag['href']
                    if href.startswith('/'):
                        href = 'https://dichvucong.danang.gov.vn' + href
                    if 'thu-tuc-hanh-chinh' in href:
                        links.append(href)

            print(f"✅ Trang {page_number}: Tìm thấy {len(links)} link")
            return crawl_detail_links(links)
        else:
            print(f"Không tải được trang {page_number}: {response.status_code}")
    except Exception as e:
        print(f"Lỗi khi truy cập trang {page_number}: {e}")
    return []

# =================== GHI RA FILE WORD ===================
def save_all_to_word(details):
    doc = Document()
    for i, item in enumerate(details, 1):
        title = item['title'].strip()
        content = item['content']

        doc.add_heading(f"THỦ TỤC {i}: {title}", level=1)

        if content.startswith("🔹 **Cơ quan thực hiện**"):
            parts = content.split("\n\n")
            if len(parts) > 1:
                parts = parts[1:] + [parts[0]]
                content = "\n\n".join(parts)

        for section in content.split("\n\n"):
            para = doc.add_paragraph()
            para.add_run(section.strip()).bold = "🔹 **" in section or "🔸 **" in section

        doc.add_paragraph("\n")

    doc.save(output_path)
    print(f"Đã lưu vào file: {output_path}")

# =================== CHẠY CHÍNH ===================
def main():
    all_details = []
    with ThreadPoolExecutor(max_workers=MAX_THREADS) as executor:
        futures = [executor.submit(crawl_page, page) for page in range(1, TOTAL_PAGES + 1)]
        for future in as_completed(futures):
            result = future.result()
            if result:
                all_details.extend(result)
    save_all_to_word(all_details)

if __name__ == "__main__":
    main()
